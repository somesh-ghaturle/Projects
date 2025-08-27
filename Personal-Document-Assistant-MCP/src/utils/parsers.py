"""
Document parsing utilities for various file formats.
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
import mimetypes

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from bs4 import BeautifulSoup
    import requests
except ImportError:
    BeautifulSoup = None
    requests = None

logger = logging.getLogger(__name__)


class DocumentParser:
    """Handles parsing of various document formats."""
    
    def __init__(self):
        """Initialize the document parser."""
        self.supported_formats = {
            '.txt': self._parse_text,
            '.md': self._parse_text,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx,
            '.html': self._parse_html,
            '.htm': self._parse_html,
        }
    
    def parse_document(self, file_path: str, encoding: str = 'utf-8') -> str:
        """Parse a document and extract text content."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_formats:
            # Try to guess MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type and mime_type.startswith('text/'):
                extension = '.txt'
            else:
                raise ValueError(f"Unsupported file format: {extension}")
        
        # Parse the document
        try:
            parser_func = self.supported_formats[extension]
            content = parser_func(file_path, encoding)
            
            if not content.strip():
                raise ValueError("Document appears to be empty")
            
            logger.info(f"Successfully parsed {extension} file: {file_path.name}")
            return content.strip()
            
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {str(e)}")
            raise
    
    def _parse_text(self, file_path: Path, encoding: str) -> str:
        """Parse plain text files."""
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for alt_encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=alt_encoding) as file:
                        logger.warning(f"Used {alt_encoding} encoding for {file_path}")
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError(f"Could not decode file with any supported encoding")
    
    def _parse_pdf(self, file_path: Path, encoding: str = None) -> str:
        """Parse PDF files."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                        continue
            
            if not text_content:
                raise ValueError("No text content found in PDF")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            raise
    
    def _parse_docx(self, file_path: Path, encoding: str = None) -> str:
        """Parse DOCX files."""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            if not text_content:
                raise ValueError("No text content found in DOCX")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise
    
    def _parse_html(self, file_path: Path, encoding: str) -> str:
        """Parse HTML files."""
        if BeautifulSoup is None:
            raise ImportError("beautifulsoup4 is required for HTML parsing. Install with: pip install beautifulsoup4")
        
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                raise ValueError("No text content found in HTML")
            
            return text
            
        except Exception as e:
            logger.error(f"HTML parsing failed: {str(e)}")
            raise
    
    def parse_url(self, url: str) -> str:
        """Parse content from a URL."""
        if requests is None or BeautifulSoup is None:
            raise ImportError("requests and beautifulsoup4 are required for URL parsing")
        
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
            
            if not text.strip():
                raise ValueError("No text content found at URL")
            
            logger.info(f"Successfully parsed content from URL: {url}")
            return text
            
        except Exception as e:
            logger.error(f"URL parsing failed for {url}: {str(e)}")
            raise
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats."""
        return list(self.supported_formats.keys())
